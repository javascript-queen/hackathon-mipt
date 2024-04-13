import os
from docx import Document
# from google.colab import files
import requests
import pandas as pd
import json
import re

# Идентификатор каталога для работы с YandexGPT
catalog_id = "b1gn3k241oqedablupl7"

# API-ключ
api_key = "AQVN2J9cCQyIVHBtyGWU6pcdhCuNoKT6gKZVz_-R"

# Роль ИИ
role = "Опытный специалист по закупкам и работе с технической документацией"

# URL YandexGPT
url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"

# Заголовки запроса
headers = {
    "Content-Type": "application/json",
    "Authorization": f"Api-Key {api_key}"
}

# # Импорт исходных данных
# def upload_files():
#     uploaded = files.upload()  # Позволяет пользователю загрузить файлы через интерфейс Colab
#     for k in uploaded.keys():
#         print('User uploaded file "{name}" with length {length} bytes'.format(
#             name=k, length=len(uploaded[k])))
#     return list(uploaded.keys())  # Возвращает список имен загруженных файлов

# # Загрузка файлов стандарта компании и предложения поставщика
# print("Пожалуйста, загрузите файл стандарта компании:")
# standard_files = upload_files()  # Может быть загружен только один файл

# print("Пожалуйста, загрузите файл предложения поставщика:")
# offer_files = upload_files()  # Может быть загружен только один файл

# # Сохранение имен файлов
# standard_file_name = standard_files[0] if standard_files else None
# offer_file_name = offer_files[0] if offer_files else None

# Функция для чтения из файла и разделение на сегменты по 9000 символов, т.к. Yandex GPT берет на вход не больше
def read_docx(file_path):
    """
    Читает документ .docx и возвращает его содержимое, разделенное на части.
    Части документа делятся так, чтобы не разрывать пункты и подпункты документа и
    начинать каждую часть с нового пункта, не превышая предел в 9000 символов.

    :param file_path: Путь к файлу .docx.
    :return: кортеж (статус, список частей содержимого или сообщение об ошибке).
    """
    if not file_path.endswith('.docx'):
        return 0, "Файл не является документом Word (.docx)"
    
    try:
        doc = Document(file_path)
        full_text = " ".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

        parts = []
        part = ""
        last_index = 0  # Индекс начала текущей части

        while last_index < len(full_text):
            # Ограничиваем текст текущей части, если возможно, до следующего пункта
            if len(full_text) - last_index > 9000:
                # Находим предел текущей части
                limit_index = last_index + 9000
                # Ищем начало следующего пункта после предела
                match = re.search(r'\s\d+\.', full_text[limit_index:])
                if match:
                    cut_index = limit_index + match.start()  # Позиция начала следующего пункта
                else:
                    cut_index = len(full_text)  # Если не найден, берем остаток текста
            else:
                cut_index = len(full_text)  # Если остаток текста меньше макс. длины, берем весь

            # Добавляем текущую часть в список
            parts.append(full_text[last_index:cut_index].strip())
            last_index = cut_index  # Обновляем индекс начала следующей части

        return 1, parts
    except Exception as e:
        return 0, f"Ошибка при чтении файла: {str(e)}"


# # Вызов функции read_docx для проверки
# status, content_or_error = read_docx(standard_file_name)

# if status == 1:
#     print("Текст Стандарта компании разделен на следующие части:")
#     for i, part in enumerate(content_or_error, 1):
#         print(f"Часть {i}: {part[:500]}...")  # Выводим первые 500 символов каждой части для просмотра
# else:
#     print(f"Ошибка: {content_or_error}")
    
# # Вызов функции read_docx для проверки
# status, content_or_error = read_docx(standard_file_name)

# if status == 1:
#     print("Текст Стандарта компании разделен на следующие части:")
#     for i, part in enumerate(content_or_error, 1):
#         print(f"Часть {i}: {part}")  # Выводим полный текст каждой части
# else:
#     print(f"Ошибка: {content_or_error}")
    
# Функция отправки запроса к Yandex GPT
def send_to_gpt(text_to_process):
    """
    Функция отправляет запрос к Yandex GPT API и возвращает результат.
    
    :param text_to_process: Текст, который будет обработан моделью GPT.
    :return: Ответ от модели GPT или сообщение об ошибке в формате словаря.
    """
    if not text_to_process:
        return {"error": "Текст для обработки не был предоставлен"}

    # Убедитесь, что catalog_id и api_key корректно определены в вашем окружении
    prompt = {
        "modelUri": f"gpt://{catalog_id}/yandexgpt-pro",
        "completionOptions": {
            "stream": False,
            "temperature": 0.6,
            "maxTokens": 2000
        },
        "messages": [
            {
                "role": "user",
                "text": text_to_process
            },
        ]
    }

    try:
        response = requests.post(url, headers=headers, json=prompt)
        response.raise_for_status()  # Это вызовет исключение, если запрос не успешен
        return response.json()  # Это должно быть словарем
    except requests.RequestException as e:
        return {"error": f"Ошибка при отправке запроса: {str(e)}"}


# Проверка
# text_to_process = "Текст, который я сейчас пришлю должен быть стандартом на оборудование (или другую продукцию) от компании-заказчика. Если это не релевантный текст - верни фразу: Этот документ не является стандартом компании. Если текст релевантный, верни только наименование оборудования, больше ничего. Если не является, то верни только это сообщение: Этот текст не является стандартом компании. Загрузите стандарт компании, чтобы я мог вам помочь сделать сравнение документов. Вот сам текст: Часть 1: Общие технические требования Содержание Введение Использование настоящего стандарта в конкретных условиях может потребовать учета дополнительных или специфических требований. Настоящий стандарт является модифицированным по отношению к международному стандарту [1] (Насосы центробежные герметичные для технологического обслуживания в нефтяной, нефтехимической и газовой промышленности). Дополнительные положения и требования, а также сноски, включенные в текст настоящего стандарта для учета потребностей национальной экономики и особенностей российской национальной стандартизации, выделены курсивом. Знак (●) в начале параграфа или его раздела указывает на то, что здесь требуется принятие решения или представление заказчиком дополнительной информации. Такую информацию необходимо привести в опросных листах или указать в запросе, или в заказе на поставку. Из соображений удобства и в информационных целях в настоящем стандарте в скобках после величин в системе единиц СИ приводятся эти же величины в системе единиц США или в других системах единиц. Область применения Настоящий стандарт устанавливает требования к герметичным центробежным насосам для использования в нефтяной, нефтехимической и газовой промышленности. Информацию о применении см. в приложении A. Примечания, входящие в приложение, носят справочный характер. Настоящий стандарт применяется к одноступенчатым насосам двух классификаций: насосам с магнитной муфтой и насосам с экранированным электродвигателем. Разделы 2-8 и раздел 10 включают требования, применимые к обеим классификациям. Раздел 9 разделен на два подраздела и включает требования, уникальные для каждой классификации. П р и м е ч а н и е — Для распространения действия стандарта на другие конструкции, такие как многоступенчатые насосы или полупогружные насосы, потребуется дополнительная информация и согласование между заказчиком и производителем. Опыт промышленной эксплуатации герметичных насосов показывает, что применение герметичных насосов, изготовленных в соответствии с требованиями настоящего стандарта, целесообразно в случаях, когда условия эксплуатации превышают любое из следующих значений: давление на выходе	1900 кПа (275 psig) давление на входе	500 кПа (75 psig) температура перекачки	150 °C (300 °F) скорость вращения	3600 об/мин номинальный общий напор	120 м (400 ft) диаметр рабочего колеса	330 мм (13”) Для насосов с уплотнениями применяется стандарт [2]. Нормативные ссылки В	настоящем	стандарте	использованы	нормативные   ссылки	на	следующие стандарты: ГОСТ 12.2.003-91	Система	стандартов	безопасности	труда.	Оборудование производственное. Общие требования безопасности ГОСТ 12.2.062-81	Система	стандартов	безопасности	труда.	Оборудование производственное. Ограждения защитные ГОСТ Р 148-1-2013 Материалы металлические. Испытание на ударный изгиб на маятниковом копре по Шарпи ГОСТ 977-88 Отливки стальные. Общие технические условия ГОСТ 3325-85 Подшипники качения. Поля допусков и технические требования к посадочным поверхностям валов и корпусов ГОСТ 6134-2007 (ИСО 9906:1999) Насосы динамические. Методы испытаний ГОСТ 6211-81 Основные нормы взаимозаменяемости. Резьба трубная коническая ГОСТ	6357-81	Основные	нормы	взаимозаменяемости.	Резьба	трубная цилиндрическая ГОСТ 7512-82 Контроль неразрушающий. Соединения сварные. Радиографический метод ГОСТ Р ИСО 8501-1-2014 Подготовка стальной поверхности перед нанесением лакокрасочных материалов и относящихся к ним продуктов. Визуальная оценка чистоты поверхности ГОСТ 8724-2002 (ИСО 261-98) Основные нормы взаимозаменяемости. Резьба метрическая. Диаметры и шаги ГОСТ 9454-78 Металлы. Метод испытания на ударный изгиб при пониженных, комнатной и повышенных температурах ГОСТ 11737-93 Ключи для винтов с внутренним шестигранником. Технические условия ГОСТ 14782-86	Контроль	неразрушающий.	Соединения	сварные.	Методы ультразвуковые. ГОСТ Р ИСО 15609-2009 Технические требования и аттестация процедур сварки металлических материалов. Технические требования к процедуре сварки ГОСТ 16093-2004	(ИСО 965-1:1998,	ИСО 965-3:1998)	Основные	нормы взаимозаменяемости. Резьба метрическая. Допуски. Посадки с зазором ГОСТ 16983-80 Ключи гаечные комбинированные. Конструкция и размеры ГОСТ 17392-72 Направляющие и опорные детали пресс-форм и форм для литья под давлением. Технические требования ГОСТ 18442-80	Контроль	неразрушающий.	Капиллярные	методы.	Общие требования ГОСТ 18855-94 (ИСО 281-89) Подшипники качения. Динамическая расчетная грузоподъемность и расчетный ресурс (долговечность) ГОСТ ИСО 1940-1-2007 Вибрация. Требования к качеству балансировки жестких роторов ГОСТ 21105-87 Контроль неразрушающий. Магнитопорошковый метод ГОСТ 23360-78 Основные нормы взаимозаменяемости. Соединения шпоночные с призматическими шпонками. Размеры шпонок и сечений пазов. Допуски и посадки ГОСТ 24069-97	(ИСО 3117-77)	Основные	нормы	взаимозаменяемости. Тангенциальные шпонки и шпоночные пазы ГОСТ 24705-2004 (ИСО 724:1993) Основные нормы взаимозаменяемости. Резьба метрическая. Основные размеры ГОСТ 24810-2013 Подшипники качения. Внутренние зазоры ГОСТ 31441.1-2011 Оборудование неэлектрическое, предназначенное для применения в потенциально взрывоопасных средах ГОСТ 28173-89 (МЭК 60034-1) Машины электрические вращающиеся. Номинальные данные и рабочие характеристики ГОСТ 28338-1989 (ИСО 6708-80) Соединения трубопроводов и арматура. Номинальные диаметры ГОСТ 31252-2004 (ИСО 3740:2000) Шум машин. Руководство по выбору метода определения уровней звуковой мощности ГОСТ 31441.1-2011	Оборудование	неэлектрическое.	Предназначенное	для применения в потенциально взрывоопасных средах. Общие требования ГОСТ 32600-2013 (ISO 21049:2004) Насосы. Уплотнительные системы вала для центробежных и роторных насосов. Общие технические требования и методы контроля ГОСТ 33259-2015 Фланцы арматуры, соединительных частей и трубопроводов на номинальное давление до PN 250. Конструкция, размеры и общие технические требования. ГОСТ Р 53687-2009 Аттестационные испытания сварщиков. Сварка плавлением. Медь и медные сплавы ГОСТ Р 53688-2009 Аттестационные испытания сварщиков. Сварка плавлением. Алюминий и алюминиевые сплавы ГОСТ Р 53690-2009 Аттестационные испытания сварщиков. Сварка плавлением. Стали ГОСТ Р 54006-2010 Аттестационные испытания сварщиков. Сварка плавлением. Никель и никелевые сплавы ГОСТ IEC 60034-1-2014 Машины электрические вращающиеся П р и м е ч а н и е – При использовании настоящего стандарта целесообразно проверить актуальность ссылочных стандартов в общедоступной информационной системе – на официальном сайте Федерального агентства по техническому регулированию и метрологии в сети Интернет или по ежегодному информационному указателю «Национальные стандарты», опубликованному по состоянию на 1 января текущего года, или по выпускам ежемесячного информационного указателя «Национальные стандарты» за текущий год. Если ссылочный стандарт заменен (изменен), то следует руководствоваться заменяющим (измененным) ссылочным стандартом. Если ссылочный стандарт отменен без замены, то положение настоящего стандарта, в котором дана ссылка на отмененный ссылочный стандарт, применяется в части, не затрагивающей эту ссылку. Все стандарты, на которые даны ссылки, в той мере, в какой они указаны в тексте, являются нормативными. Примечания, следующие за разделами, носят справочный характер. Стандарты, коды и технические условия, действующие на момент публикации настоящего стандарта, в той мере, в какой это указано в настоящем документе, являются частью настоящего стандарта. Применимость изменений в стандартах, кодексах и технических условиях, произошедших после публикации настоящего стандарта, должна быть взаимно согласована между заказчиком и производителем. Термины и определения В настоящем стандарте применены следующие термины с соответствующими определениями: воздушный зазор (air gap): Радиальное расстояние между внутренней стороной внешнего магнитного кольца и внешней стороной защитной оболочки. корпус с осевым разъемом (axially split): Корпус насоса, главный разъем которого расположен параллельно оси вала. допустимая рабочая область (диапазон) (allowable operating region): Диапазон основных параметров насоса, при работе внутри которого не достигаются предельные значения вибрации или температуры, устанавливаемые в настоящем стандарте, или другие предельные значения. установленные производителем. осевая сила (axial thrust): Общая осевая нагрузка на вал насоса, вызванная воздействием гидравлической силы на лопатки рабочего колеса, ротор и на внутреннюю магнитную полумуфту. точка наивысшего КПД; BEР (best efficiency point): Рабочая точка (подача- напор), в которой насос имеет максимальный коэффициент полезного действия КПД при расчетном диаметре рабочего колеса. П р и м е ч а н е"
# print (send_to_gpt(text_to_process))
    
# Извлечение наименование оборудования
def extract_equipment_name(introduction):
    """
    Функция для извлечения наименования оборудования из стандартного документа с минимальным запросом.
    """

    minimal_query = introduction
    minimal_query = f"Текст, который я сейчас пришлю должен быть стандартом на оборудование (или другую продукцию) от компании-заказчика. Если это не релевантный текст - верни фразу: Этот документ не является стандартом компании. Если текст релевантный, верни только наименование оборудования, больше ничего. Если не является, то верни 0. Вот сам текст: {introduction}"
    equipment_name_result = send_to_gpt(minimal_query)

    if 'error' in equipment_name_result:
        return {"error": equipment_name_result['error']}

    equipment_name = (
        equipment_name_result
            .get('result', {})
            .get('alternatives', {})[0]
            .get('message', {})
            .get('text', '')
            .strip()
    )
    if equipment_name == 0:
        return {"error": "Не удалось извлечь наименование оборудования из документа."}
    if not equipment_name:
        return {"error": "Не удалось извлечь наименование оборудования из документа."}

    return {"equipment_name": equipment_name}

# # Функция анализа и разделения фрагментов на смысловые блоки   
def analyze_document(file_path):
    status, parts = read_docx(file_path)
    if status == 0:
        return pd.DataFrame(), f"Ошибка при чтении файла: {parts}"

    sections_list = []
    for part in parts:
        section_request_text = f"Раздели текст на смысловые разделы и выдай списком тексты каждого раздела. Вот текст: {part}"
        section_result = send_to_gpt(section_request_text)
        if "error" in section_result:
            return pd.DataFrame(), section_result['error']
        
        # Проверяем, что ответ содержит 'choices' и что в 'choices' есть элементы
        if 'choices' in section_result and section_result['choices']:
            sections = section_result['choices'][0]['text'].split("\n")
            sections_list.extend(sections)
        else:
            return pd.DataFrame(), "API не вернул ожидаемые данные."

    return pd.DataFrame(sections_list, columns=['Section']), "Success"
 
 
 # Функция анализа и разделения фрагментов на смысловые блоки без GPT   
def create_dataframe_from_docx(file_path):
    # Чтение документа
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Объединение текста в одну строку для упрощения разбиения
    full_text = " ".join(full_text)

    # Регулярное выражение для выделения секций
    # \d+ - одна или более цифр (номер раздела)
    # (\.\d+)* - ноль или более последовательностей из точки и одной или более цифр (подразделы)
    # \s - пробел
    # [^\d] - любой символ, кроме цифры (чтобы исключить ситуации, где за номером секции следует другой номер)
    pattern = r"(\d+(\.\d+)*\.\s[^\d])"

    # Находим все совпадения и добавляем позиции начала совпадений в список
    matches = [match.start() for match in re.finditer(pattern, full_text)]

    # Создаем список секций, используя найденные позиции для разбиения текста
    sections = [full_text[matches[i]:matches[i + 1]] for i in range(len(matches) - 1)]
    sections.append(full_text[matches[-1]:])  # Добавляем последний фрагмент

    # Создание DataFrame
    df = pd.DataFrame({'Section': sections})
    return df



def filter_and_extract(sections_df, equipment_name, gpt_prompt):
    """
    Фильтрует и извлекает характеристики оборудования из DataFrame смысловых разделов.

    :param sections_df: DataFrame смысловых разделов.
    :param equipment_name: Наименование оборудования.
    :param gpt_prompt: Запрос к нейросети.
    :return: DataFrame с характеристиками и сообщение об успешности операции.
    """
    # Запрашиваем ответ у нейросети
    characteristic_check = f"{gpt_prompt} '{equipment_name}'? Ответь на вопрос по каждому пункту только да или нет после каждого !конец пункта!. Отвечай да только при наличии конкретных характеристик оборудования. "
    check_result = send_to_gpt(characteristic_check)
    
    # Проверяем наличие ошибки при обращении к нейросети
    if "error" in check_result:
        return pd.DataFrame(), check_result['error']
    
    # Проверяем ответ от нейросети
    response_text = check_result['result']['alternatives'][0]['message']['text'].strip().lower()
    
    # Инициализируем список для характеристик
    characteristics = []
    
    # Обрабатываем каждый раздел
    for index, row in sections_df.iterrows():
        # Добавляем маркер в конец каждого пункта
        section_text = row['Section'] + " !конец пункта! "
        
        # Разбиваем текст на пункты
        extracted_characteristics = re.split(r'\d+\.\s+', section_text)
        
        # Добавляем каждый пункт с ответом в список, если ответ "да"
        for characteristic in extracted_characteristics:
            characteristic = characteristic.strip()
            if characteristic and response_text == 'да':
                characteristics.append({'section': characteristic, 'answer': response_text})
    
    # Создаем DataFrame из списка характеристик
    characteristics_df = pd.DataFrame(characteristics)
    
    return characteristics_df, "Success"


# # Функция анализа и разделения фрагментов на смысловые блоки (старый вариант)  
 
# def analyze_document(file_path):
#     status, parts = read_docx(file_path)
#     if status == 0:
#         return pd.DataFrame(), f"Ошибка при чтении файла: {parts}"

#     sections_list = []
#     for part in parts:
#         section_request_text = f"Раздели текст на смысловые разделы и выдай списком тексты каждого раздела. Вот текст: {part}"
#         section_result = send_to_gpt(section_request_text)
#         if "error" in section_result:
#             return pd.DataFrame(), section_result['error']
        
#         sections = section_result['choices'][0]['text'].split("\n")
#         sections_list.extend(sections)

    # return pd.DataFrame(sections_list, columns=['Section']), "Success"

# Функция выделения ключевых характеристик
import re
import pandas as pd

def filter_and_extract(sections_df, equipment_name, gpt_prompt):
    """
    Фильтрует и извлекает характеристики оборудования из DataFrame смысловых разделов.

    :param sections_df: DataFrame смысловых разделов.
    :param equipment_name: Наименование оборудования.
    :param gpt_prompt: Запрос к нейросети.
    :return: DataFrame с характеристиками и сообщение об успешности операции.
    """
    # Запрашиваем ответ у нейросети
    characteristic_check = f"{gpt_prompt} '{equipment_name}'? Ответь на вопрос по каждому пункту только да или нет после каждого !конец пункта!. Отвечай да только при наличии конкретных характеристик оборудования. "
    check_result = send_to_gpt(characteristic_check)
    
    # Проверяем наличие ошибки при обращении к нейросети
    if "error" in check_result:
        return pd.DataFrame(), check_result['error']
    
    # Проверяем ответ от нейросети
    response_text = check_result['result']['alternatives'][0]['message']['text'].strip().lower()
    
    # Инициализируем список для характеристик
    characteristics = []
    
    # Обрабатываем каждый раздел
    for index, row in sections_df.iterrows():
        # Добавляем маркер в конец каждого пункта
        section_text = row['Section'] + " !конец пункта! "
        
        # Разбиваем текст на пункты
        extracted_characteristics = re.split(r'\d+\.\s+', section_text)
        
        # Добавляем каждый пункт с ответом в список, если ответ "да"
        for characteristic in extracted_characteristics:
            characteristic = characteristic.strip()
            if characteristic and response_text == 'да':
                characteristics.append({'section': characteristic, 'answer': response_text})
    
    # Создаем DataFrame из списка характеристик
    characteristics_df = pd.DataFrame(characteristics)
    
    return characteristics_df, "Success"



# Функция сравнения смысловых блоков в двух списках ключевых характеристик
def compare_characteristics(standard_df, offer_df):
    """
    Сравнивает характеристики оборудования между стандартом и предложением.
    
    :param standard_df: DataFrame с характеристиками из стандарта.
    :param offer_df: DataFrame с характеристиками из предложения.
    :return: DataFrame с результатами сравнения.
    """
    # Создаем итоговый DataFrame
    comparison_df = pd.DataFrame(columns=['Characteristic', 'Standard', 'Offer', 'Match'])
    
    # Преобразуем DataFrame для удобства сравнения
    standard_df.set_index('Characteristic', inplace=True)
    offer_df.set_index('Characteristic', inplace=True)
    
    # Объединяем оба DataFrame для сравнения
    all_characteristics = standard_df.join(offer_df, lsuffix='_std', rsuffix='_off', how='outer')
    
    # Заполняем итоговый DataFrame
    for characteristic in all_characteristics.index:
        standard_val = all_characteristics.loc[characteristic, 'Value_std'] if 'Value_std' in all_characteristics.columns else None
        offer_val = all_characteristics.loc[characteristic, 'Value_off'] if 'Value_off' in all_characteristics.columns else None
        is_match = 'Yes' if standard_val == offer_val else 'No'
        
        comparison_df = comparison_df.append({
            'Characteristic': characteristic,
            'Standard': standard_val,
            'Offer': offer_val,
            'Match': is_match
        }, ignore_index=True)

    return comparison_df

def process_documents(standard_file_path, offer_file_path):
    """
    Основная функция для обработки документов стандарта и предложения.
    """
    if not standard_file_path or not offer_file_path:
        missing_files = " и ".join([f"файл стандарта компании" if not standard_file_path else "",
                                    f"файл предложения от поставщика" if not offer_file_path else ""]).strip()
        return {"error": f"Не загружены: {missing_files}"}

    standard_status, standard_parts = read_docx(standard_file_path)
    if standard_status == 0:
        return {"error": f"Ошибка при чтении файла стандарта компании: {standard_parts}"}

    offer_status, offer_parts = read_docx(offer_file_path)
    if offer_status == 0:
        return {"error": f"Ошибка при чтении файла предложения от поставщика: {offer_parts}"}

    # Extract equipment name form Company Standards
    equipment_name = extract_equipment_name(standard_parts[0])
    # print (equipment_name)

    # Создаем DataFrame из стандартного документа
    standard_df = create_dataframe_from_docx(standard_file_path)

    # Создаем DataFrame из документа предложения от поставщика
    offer_df = create_dataframe_from_docx(offer_file_path)

    # Фильтруем характеристики в стандарте
    standard_characteristics_df, standard_message = filter_and_extract(standard_df, equipment_name,
                                                                             "Проанализируй текст и раздели его на полные смысловые разделы, соответствующие отдельным подпунктам (название и содержание пропусти, если они есть). Каждый раздел должен начинаться с номера пункта или подпункта (обязательно!) (например, 1, 1.1, 1.2.4.7 и т.д.) и содержать полный текст одного пункта или подпункта, заканчивая последним предложением перед следующим пунктом или подпунктом.")
    
    # Извлекаем характеристики из предложения
    offer_characteristics_df, offer_message = filter_and_extract(offer_df, equipment_name,
                                                                 "Содержит ли каждый из пунктов конкретные характеристики данного оборудования")

    # Возвращаем результаты
    return {
        "standard_df": standard_df,
        "offer_df": offer_df,
        "standard_characteristics_df": standard_characteristics_df,
        "offer_characteristics_df": offer_characteristics_df,
        "message": "Документы успешно обработаны и разделены на секции.",
        "standard_message": standard_message,
        "offer_message": offer_message
    }

# Вызов основной функции
result = process_documents(standard_file_name, offer_file_name)
if "error" not in result:
    print("Стандартный документ разделы:")
    print(result["standard_df"].head())
    print("\nПредложение от поставщика разделы:")
    print(result["offer_df"].head())
    print("\nХарактеристики из стандарта:")
    print(result["standard_characteristics_df"].head())
    print("\nХарактеристики из предложения:")
    print(result["offer_characteristics_df"].head())
    print("\nСообщение о стандарте:", result["standard_message"])
    print("Сообщение о предложении:", result["offer_message"])
else:
    print(result["error"])

# Вариации основной функции 
# def process_documents(standard_file_path, offer_file_path):
#     """
#     Основная функция для обработки документов стандарта и предложения.
#     """
#     if not standard_file_path or not offer_file_path:
#         missing_files = " и ".join([f"файл стандарта компании" if not standard_file_path else "",
#                                     f"файл предложения от поставщика" if not offer_file_path else ""]).strip()
#         return {"error": f"Не загружены: {missing_files}"}
    
#     standard_status, standard_parts = read_docx(standard_file_path)
#     if standard_status == 0:
#         return {"error": f"Ошибка при чтении файла стандарта компании: {standard_parts}"}
    
#     offer_status, offer_parts = read_docx(offer_file_path)
#     if offer_status == 0:
#         return {"error": f"Ошибка при чтении файла предложения от поставщика: {offer_parts}"}

#     if standard_parts:
#         first_part = standard_parts[0]
#         request_text = f"Текст, который я сейчас пришлю должен быть стандартом на оборудование (или другую продукцию) от компании-заказчика. Если это не релевантный текст - верни 0. Вот сам текст: {first_part}"
#         gpt_result = send_to_gpt(request_text)
#         if not gpt_result or "error" in gpt_result:
#             return gpt_result
#         return gpt_result    
# Этот текст не является стандартом компании. Загрузите стандарт компании, чтобы я мог вам помочь сделать сравнение документов.

#     # if standard_parts:
#     #     first_part = standard_parts[0]
#     #     request_text = f"Раздели текст на смысловые разделы и выдай текст каждого раздела, особенно интересует 'Базовая конструкция': {first_part}"
#     #     gpt_result = send_to_gpt(request_text)
#     #     if not gpt_result or "error" in gpt_result:
#     #         return gpt_result
#     #     return gpt_result

#     return {"error": "Неизвестная ошибка при обработке документов"}

# # Предполагается, что вы где-то определили standard_file_name и offer_file_name
# # result = process_documents(standard_file_name, offer_file_name)
# # print(result)



# # Вызов основной функции
# result = process_documents(standard_file_name, offer_file_name)
# print(result)
    

# def process_documents(standard_file_path, offer_file_path):
#     """
#     Основная функция для обработки документов стандарта и предложения.
#     """
#     if not standard_file_path or not offer_file_path:
#         missing_files = " и ".join([f"файл стандарта компании" if not standard_file_path else "",
#                                     f"файл предложения от поставщика" if not offer_file_path else ""]).strip()
#         return {"error": f"Не загружены: {missing_files}"}
    
#     standard_status, standard_parts = read_docx(standard_file_path)
#     if standard_status == 0:
#         return {"error": f"Ошибка при чтении файла стандарта компании: {standard_parts}"}
    
#     offer_status, offer_parts = read_docx(offer_file_path)
#     if offer_status == 0:
#         return {"error": f"Ошибка при чтении файла предложения от поставщика: {offer_parts}"}

#     # ЗАГЛУШКА
#     equipment_name = "Насосы центробежные герметичные для технологического обслуживания в нефтяной, нефтехимической и газовой промышленности"

#     # Извлечение и фильтрация разделов из стандарта
#     standard_sections = extract_relevant_sections(standard_parts, equipment_name)
#     if "error" in standard_sections:
#         return standard_sections

#     # Извлечение и фильтрация разделов из предложения
#     offer_sections = extract_relevant_sections(offer_parts, equipment_name)
#     if "error" in offer_sections:
#         return offer_sections

#     return {"standard_sections": standard_sections, "offer_sections": offer_sections}

# def extract_relevant_sections(document_parts, equipment_name):
#     relevant_sections = []
#     for part in document_parts:
#         section_request_text = f"Раздели текст на смысловые разделы и выдай списком тексты каждого раздела. Вот текст: {part[:1000]}"  # Ограничиваем текст для упрощения
#         section_result = send_to_gpt(section_request_text)
#         if "error" in section_result:
#             return {"error": f"Ошибка при обработке разделов: {section_result['error']}"}
        
#         for section in section_result.split("\n"):
#             relevance_request_text = f"Относится ли этот пункт к {equipment_name}? Текст пункта: {section}"
#             relevance_result = send_to_gpt(relevance_request_text)
#             if "error" in relevance_result:
#                 return {"error": f"Ошибка

# Функция для тестирования - работает
# # def read_docx(file_name):
# #     doc = Document(file_name)
# #     full_text = " ".join([paragraph.text for paragraph in doc.paragraphs])
# #     return full_text[:10000] # Ограничим текст для упрощения



# # file_name = 'pump_standards.docx'
# # text_to_process = read_docx(file_name)
# # text_to_process


# # Текст для обработки моделью Yandex GPT
# prompt = {
#     "modelUri": f"gpt://{catalog_id}/yandexgpt-pro",
#     "completionOptions": {
#         "stream": False,
#         "temperature": 0.6,
#         "maxTokens": "2000"
#     },
#     "messages": [
#         {
#             "role": "user",
#             "text": f"Привет, можешь разделить исходный текст на смысловые разделы и выдать текст каждого раздела по порядку: {text_to_process}"
#         },
#     ]
# }


# # Отправляем POST-запрос к Yandex GPT API
# response = requests.post(url, headers=headers, json=prompt)

# # Выводим результат запроса
# print(response.json())

# print(response.json()['result']['alternatives'][0]['message']['text'])